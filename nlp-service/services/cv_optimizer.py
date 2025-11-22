import os
import shutil
import re
from typing import Dict, List
import logging
from docx import Document
import PyPDF2

logger = logging.getLogger(__name__)


class CVOptimizer:
    """
    Personnalise un CV pour une offre d'emploi spécifique
    """
    
    def __init__(self):
        self.output_dir = os.getenv('CV_OUTPUT_DIR', './optimized_cvs')
        os.makedirs(self.output_dir, exist_ok=True)
    
    def optimize(self, cv_path: str, job_description: str, 
                 job_requirements: str, job_title: str) -> Dict:
        """
        Optimise un CV pour une offre d'emploi
        """
        try:
            # Lire le CV original
            file_ext = os.path.splitext(cv_path)[1].lower()
            
            if file_ext == '.pdf':
                # Pour les PDFs, on crée une version Word optimisée
                text = self._extract_from_pdf(cv_path)
                optimized_path = self._optimize_text_cv(cv_path, text, job_description, job_requirements, job_title)
            elif file_ext in ['.doc', '.docx']:
                optimized_path = self._optimize_docx_cv(cv_path, job_description, job_requirements, job_title)
            else:
                raise ValueError(f'Unsupported file format: {file_ext}')
            
            # Analyser les changements
            changes = self._analyze_changes(cv_path, optimized_path, job_description, job_requirements)
            
            return {
                'customized_path': optimized_path,
                'changes': changes,
                'match_improvements': self._calculate_improvements(changes)
            }
        except Exception as e:
            logger.error(f'Error optimizing CV: {str(e)}')
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extrait le texte d'un PDF"""
        text = ''
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + '\n'
        except Exception as e:
            logger.error(f'Error extracting PDF text: {str(e)}')
            raise
        return text
    
    def _optimize_text_cv(self, original_path: str, text: str, 
                          job_description: str, job_requirements: str, job_title: str) -> str:
        """
        Optimise un CV texte (pour les PDFs)
        """
        # Créer un nouveau document Word optimisé
        doc = Document()
        
        # Extraire les mots-clés de l'offre
        job_keywords = self._extract_keywords(job_description + ' ' + job_requirements)
        required_skills = self._extract_skills(job_requirements)
        
        # Optimiser le texte
        optimized_text = self._add_keywords_to_text(text, job_keywords, required_skills)
        
        # Ajouter au document
        for paragraph in optimized_text.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        # Sauvegarder
        base_name = os.path.splitext(os.path.basename(original_path))[0]
        output_path = os.path.join(self.output_dir, f'{base_name}_optimized.docx')
        doc.save(output_path)
        
        return output_path
    
    def _optimize_docx_cv(self, cv_path: str, job_description: str, 
                         job_requirements: str, job_title: str) -> str:
        """
        Optimise un CV Word
        """
        doc = Document(cv_path)
        
        # Extraire les mots-clés
        job_keywords = self._extract_keywords(job_description + ' ' + job_requirements)
        required_skills = self._extract_skills(job_requirements)
        
        # Optimiser chaque paragraphe
        for paragraph in doc.paragraphs:
            if paragraph.text:
                optimized_text = self._add_keywords_to_text(
                    paragraph.text, 
                    job_keywords, 
                    required_skills
                )
                paragraph.text = optimized_text
        
        # Sauvegarder
        base_name = os.path.splitext(os.path.basename(cv_path))[0]
        output_path = os.path.join(self.output_dir, f'{base_name}_optimized.docx')
        doc.save(output_path)
        
        return output_path
    
    def _extract_keywords(self, text: str) -> List[str]:
        """
        Extrait les mots-clés importants d'un texte
        """
        stop_words = {'le', 'la', 'les', 'un', 'une', 'des', 'de', 'du', 'et', 'ou', 'à', 'pour', 'avec'}
        words = re.findall(r'\b\w{4,}\b', text.lower())
        keywords = [w for w in words if w not in stop_words]
        
        # Compter les occurrences
        from collections import Counter
        keyword_counts = Counter(keywords)
        
        # Retourner les 20 mots-clés les plus fréquents
        return [word for word, count in keyword_counts.most_common(20)]
    
    def _extract_skills(self, text: str) -> List[str]:
        """
        Extrait les compétences mentionnées
        """
        common_skills = [
            'python', 'java', 'javascript', 'typescript', 'react', 'vue', 'angular',
            'node.js', 'django', 'flask', 'spring', 'postgresql', 'mysql', 'mongodb',
            'aws', 'azure', 'docker', 'kubernetes', 'git', 'agile', 'scrum',
            'machine learning', 'deep learning', 'data science', 'big data'
        ]
        
        text_lower = text.lower()
        found_skills = [skill for skill in common_skills if skill in text_lower]
        
        return found_skills
    
    def _add_keywords_to_text(self, text: str, keywords: List[str], skills: List[str]) -> str:
        """
        Ajoute des mots-clés pertinents au texte sans le dénaturer
        """
        optimized = text
        
        # Ajouter des compétences manquantes dans la section compétences
        if 'compétence' in text.lower() or 'skill' in text.lower():
            # Trouver la section compétences et ajouter les compétences manquantes
            skills_in_text = [s for s in skills if s in text.lower()]
            missing_skills = [s for s in skills if s not in text.lower()]
            
            if missing_skills:
                # Ajouter les compétences manquantes
                optimized = optimized + '\n' + ', '.join(missing_skills[:5])  # Limiter à 5
    
        return optimized
    
    def _analyze_changes(self, original_path: str, optimized_path: str, 
                        job_description: str, job_requirements: str) -> List[Dict]:
        """
        Analyse les changements apportés au CV
        """
        changes = []
        
        # Lire les deux fichiers
        original_text = self._read_file(original_path)
        optimized_text = self._read_file(optimized_path)
        
        # Extraire les mots-clés de l'offre
        job_keywords = self._extract_keywords(job_description + ' ' + job_requirements)
        
        # Identifier les mots-clés ajoutés
        original_lower = original_text.lower()
        optimized_lower = optimized_text.lower()
        
        added_keywords = []
        for keyword in job_keywords:
            if keyword not in original_lower and keyword in optimized_lower:
                added_keywords.append(keyword)
        
        if added_keywords:
            changes.append({
                'type': 'keywords_added',
                'description': f'Ajout de {len(added_keywords)} mots-clés pertinents',
                'details': added_keywords[:10]  # Limiter à 10
            })
        
        # Identifier les compétences ajoutées
        required_skills = self._extract_skills(job_requirements)
        added_skills = [s for s in required_skills if s not in original_lower and s in optimized_lower]
        
        if added_skills:
            changes.append({
                'type': 'skills_added',
                'description': f'Ajout de {len(added_skills)} compétences',
                'details': added_skills
            })
        
        return changes
    
    def _read_file(self, file_path: str) -> str:
        """
        Lit le contenu d'un fichier
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_ext in ['.doc', '.docx']:
            doc = Document(file_path)
            return '\n'.join([p.text for p in doc.paragraphs])
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
    
    def _calculate_improvements(self, changes: List[Dict]) -> Dict:
        """
        Calcule l'amélioration du score de matching
        """
        improvements = {
            'keywords_score_improvement': 0,
            'skills_score_improvement': 0,
            'total_improvement': 0
        }
        
        for change in changes:
            if change['type'] == 'keywords_added':
                improvements['keywords_score_improvement'] += len(change['details']) * 2
            elif change['type'] == 'skills_added':
                improvements['skills_score_improvement'] += len(change['details']) * 5
        
        improvements['total_improvement'] = (
            improvements['keywords_score_improvement'] + 
            improvements['skills_score_improvement']
        )
        
        return improvements

