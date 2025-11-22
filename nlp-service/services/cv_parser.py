import os
import re
import json
from typing import Dict, List, Optional
import PyPDF2
from docx import Document
import logging

logger = logging.getLogger(__name__)

# Import OpenAI parser si disponible
try:
    from services.openai_cv_parser import OpenAICVParser
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    logger.warning('OpenAI parser not available. Using fallback parser.')


class CVParser:
    """
    Parse les CVs (PDF, DOC, DOCX) et extrait les informations structurées
    Utilise OpenAI en priorité si disponible, sinon utilise l'extraction regex
    """
    
    def __init__(self):
        self.skills_keywords = self._load_skills_keywords()
        self.openai_parser = None
        
        # Initialiser OpenAI parser si disponible
        if OPENAI_AVAILABLE:
            try:
                self.openai_parser = OpenAICVParser()
                logger.info('OpenAI CV parser initialized successfully')
            except Exception as e:
                logger.warning(f'OpenAI parser initialization failed: {str(e)}. Using fallback parser.')
                self.openai_parser = None
        
        # Initialiser Spacy
        try:
            import spacy
            try:
                self.nlp_fr = spacy.load("fr_core_news_sm")
                self.nlp_en = spacy.load("en_core_web_sm")
                self.spacy_available = True
                logger.info('Spacy models loaded successfully')
            except OSError:
                logger.warning('Spacy models not found. Downloading...')
                from spacy.cli import download
                download("fr_core_news_sm")
                download("en_core_web_sm")
                self.nlp_fr = spacy.load("fr_core_news_sm")
                self.nlp_en = spacy.load("en_core_web_sm")
                self.spacy_available = True
        except Exception as e:
            logger.warning(f'Spacy initialization failed: {str(e)}')
            self.spacy_available = False

    def parse(self, file_path: str) -> Dict:
        """
        Parse un fichier CV et retourne les données structurées
        """
        # Convertir en chemin absolu si nécessaire
        if not os.path.isabs(file_path):
            # Si c'est un chemin relatif, essayer depuis le répertoire backend/uploads
            possible_paths = [
                file_path,
                os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'backend', 'uploads', os.path.basename(file_path)),
                os.path.join(os.getcwd(), file_path),
            ]
            file_path = None
            for path in possible_paths:
                if os.path.exists(path):
                    file_path = os.path.abspath(path)
                    break
            
            if not file_path or not os.path.exists(file_path):
                raise FileNotFoundError(f'CV file not found. Tried: {possible_paths}')
        else:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f'CV file not found: {file_path}')
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                text = self._extract_from_pdf(file_path)
            elif file_ext in ['.doc', '.docx']:
                text = self._extract_from_docx(file_path)
            else:
                raise ValueError(f'Unsupported file format: {file_ext}')
            
            # Vérifier que du texte a été extrait
            if not text or len(text.strip()) < 10:
                raise ValueError('No text could be extracted from the CV file. The file might be corrupted or image-based.')
            
            # Utiliser OpenAI si disponible
            if self.openai_parser:
                try:
                    logger.info('Using OpenAI to parse CV')
                    parsed_data = self.openai_parser.parse_from_text(text)
                    logger.info('CV parsed successfully with OpenAI')
                    return parsed_data
                except Exception as e:
                    logger.warning(f'OpenAI parsing failed: {str(e)}. Falling back to Spacy/Regex.')
            
            # Fallback: Spacy + Regex
            logger.info('Using Spacy/Regex extraction for CV parsing')
            
            # Déterminer la langue (simplifié)
            is_french = 'français' in text.lower() or 'expérience' in text.lower()
            nlp = self.nlp_fr if self.spacy_available and is_french else (self.nlp_en if self.spacy_available else None)
            
            spacy_data = {}
            if nlp:
                spacy_data = self._extract_with_spacy(text, nlp)
            
            # Combiner Spacy et Regex (Regex remplit les trous)
            parsed_data = {
                'raw_text': text,
                'personal_info': {**self._extract_personal_info(text), **spacy_data.get('personal_info', {})},
                'skills': list(set(self._extract_skills(text) + spacy_data.get('skills', []))),
                'experience': spacy_data.get('experience') or self._extract_experience(text),
                'education': spacy_data.get('education') or self._extract_education(text),
                'languages': self._extract_languages(text),
                'certifications': self._extract_certifications(text),
                'projects': self._extract_projects(text),
                'summary': self._extract_summary(text),
            }
            
            return parsed_data
        except Exception as e:
            logger.error(f'Error parsing CV: {str(e)}', exc_info=True)
            raise

    def _extract_with_spacy(self, text: str, nlp) -> Dict:
        """Extrait les informations avec Spacy NER"""
        doc = nlp(text[:100000]) # Limite de taille pour Spacy
        
        data = {
            'personal_info': {},
            'skills': [],
            'experience': [],
            'education': []
        }
        
        # Extraction d'entités
        for ent in doc.ents:
            # Personne (Nom)
            if ent.label_ == "PER" and not data['personal_info'].get('full_name'):
                data['personal_info']['full_name'] = ent.text
            
            # Organisation (Entreprise / École)
            elif ent.label_ == "ORG":
                # Heuristique simple: si proche de mots clés "éducation", c'est une école
                # Sinon entreprise
                pass # Difficile de distinguer sans contexte, on laisse Regex faire le gros du travail structuré
            
            # Localisation
            elif ent.label_ == "LOC" and not data['personal_info'].get('location'):
                data['personal_info']['location'] = ent.text
        
        return data
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extrait le texte d'un PDF"""
        text = ''
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Vérifier que le PDF n'est pas vide
                if len(pdf_reader.pages) == 0:
                    raise ValueError('PDF file is empty or corrupted')
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + '\n'
                        else:
                            logger.warning(f'No text extracted from page {page_num + 1}')
                    except Exception as page_error:
                        logger.warning(f'Error extracting text from page {page_num + 1}: {str(page_error)}')
                        continue
                
                if not text.strip():
                    raise ValueError('No text could be extracted from PDF. The PDF might be image-based or encrypted.')
                    
        except Exception as e:
            logger.error(f'Error extracting PDF text: {str(e)}')
            raise ValueError(f'Failed to extract text from PDF: {str(e)}')
        return text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extrait le texte d'un document Word"""
        try:
            doc = Document(file_path)
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extraire aussi les tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs.append(row_text)
            
            text = '\n'.join(paragraphs)
            
            if not text.strip():
                raise ValueError('No text could be extracted from DOCX file')
            
            return text
        except Exception as e:
            logger.error(f'Error extracting DOCX text: {str(e)}')
            raise ValueError(f'Failed to extract text from DOCX: {str(e)}')
    
    def _extract_personal_info(self, text: str) -> Dict:
        """Extrait les informations personnelles"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'(\+33|0)[1-9](\d{2}){4}'
        
        emails = re.findall(email_pattern, text)
        phones = re.findall(phone_pattern, text)
        
        # Tentative d'extraction du nom (très basique: première ligne ou lignes avec mots clés)
        # C'est difficile sans NLP, donc on laisse souvent vide ou on prend le début
        
        return {
            'full_name': None, # Difficile à extraire avec regex fiable
            'email': emails[0] if emails else None,
            'phone': phones[0] if phones else None,
            'location': None,
            'linkedin': None,
            'website': None
        }
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extrait les compétences techniques"""
        text_lower = text.lower()
        found_skills = []
        
        # Recherche de compétences communes
        for skill_category, skills in self.skills_keywords.items():
            for skill in skills:
                if skill.lower() in text_lower:
                    found_skills.append(skill)
        
        # Recherche de sections "Compétences" ou "Skills"
        skills_section_pattern = r'(?:compétences?|skills?|technologies?)[\s:]*\n(.*?)(?:\n\n|\n[A-Z])'
        match = re.search(skills_section_pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            skills_text = match.group(1)
            # Extraire les compétences de cette section
            skills_list = [s.strip() for s in re.split(r'[,;•\-\n]', skills_text) if s.strip()]
            found_skills.extend(skills_list)
        
        return list(set(found_skills))  # Supprimer les doublons
    
    def _extract_experience(self, text: str) -> List[Dict]:
        """Extrait les expériences professionnelles"""
        experiences = []
        
        # Pattern pour détecter les expériences
        # Format: Poste - Entreprise (Date - Date)
        exp_pattern = r'([A-Z][^•\n]+?)\s*[-–]\s*([A-Z][^•\n]+?)\s*\(([^)]+)\)'
        
        matches = re.finditer(exp_pattern, text)
        for match in matches:
            experiences.append({
                'position': match.group(1).strip(),
                'company': match.group(2).strip(),
                'start_date': match.group(3).split('-')[0].strip() if '-' in match.group(3) else match.group(3).strip(),
                'end_date': match.group(3).split('-')[1].strip() if '-' in match.group(3) else None,
                'description': None,
                'location': None
            })
        
        return experiences
    
    def _extract_education(self, text: str) -> List[Dict]:
        """Extrait les formations"""
        education = []
        
        # Mots-clés pour identifier les sections éducation
        education_keywords = ['diplôme', 'formation', 'éducation', 'education', 'degree', 'university', 'école', 'school']
        
        # Pattern pour détecter les formations
        edu_pattern = r'([A-Z][^•\n]+?)\s*[-–]\s*([A-Z][^•\n]+?)\s*\(([^)]+)\)'
        
        # Rechercher dans les sections pertinentes
        lines = text.split('\n')
        in_education_section = False
        
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in education_keywords):
                in_education_section = True
                continue
            
            if in_education_section:
                match = re.search(edu_pattern, line)
                if match:
                    education.append({
                        'degree': match.group(1).strip(),
                        'institution': match.group(2).strip(),
                        'start_date': match.group(3).split('-')[0].strip() if '-' in match.group(3) else match.group(3).strip(),
                        'end_date': match.group(3).split('-')[1].strip() if '-' in match.group(3) else None,
                        'field_of_study': None,
                        'location': None
                    })
        
        return education
    
    def _extract_languages(self, text: str) -> List[Dict]:
        """Extrait les langues"""
        languages = []
        common_languages = ['français', 'anglais', 'espagnol', 'allemand', 'italien', 
                           'french', 'english', 'spanish', 'german', 'italian']
        
        text_lower = text.lower()
        for lang in common_languages:
            if lang in text_lower:
                languages.append({
                    'language': lang.capitalize(),
                    'level': None
                })
        
        return languages
    
    def _extract_certifications(self, text: str) -> List[Dict]:
        """Extrait les certifications"""
        certifications = []
        
        cert_keywords = ['certification', 'certificat', 'certified', 'certificate']
        cert_pattern = r'(?:certification|certificat|certified|certificate)[\s:]*([A-Z][^•\n]+)'
        
        matches = re.finditer(cert_pattern, text, re.IGNORECASE)
        for match in matches:
            certifications.append({
                'name': match.group(1).strip(),
                'issuer': None,
                'date': None
            })
        
        return certifications

    def _extract_projects(self, text: str) -> List[Dict]:
        """Extrait les projets (Regex fallback)"""
        projects = []
        # Recherche basique de section Projets
        project_section_pattern = r'(?:projets?|projects?)[\s:]*\n(.*?)(?:\n\n|\n[A-Z])'
        match = re.search(project_section_pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            project_text = match.group(1)
            # Essayer de séparer par lignes ou puces
            items = [p.strip() for p in re.split(r'[\n•-]', project_text) if p.strip()]
            for item in items:
                projects.append({
                    'name': item,
                    'description': None,
                    'technologies': [],
                    'url': None
                })
        return projects

    def _extract_summary(self, text: str) -> Optional[str]:
        """Extrait le résumé professionnel (Regex fallback)"""
        # Prend souvent le premier paragraphe significatif ou section "Profil"
        summary_pattern = r'(?:profil|profile|summary|résumé|objectif)[\s:]*\n(.*?)(?:\n\n|\n[A-Z])'
        match = re.search(summary_pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            return match.group(1).strip()
        return None
    
    def _load_skills_keywords(self) -> Dict[str, List[str]]:
        """Charge les mots-clés de compétences"""
        return {
            'programming_languages': [
                'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'Go', 'Rust',
                'PHP', 'Ruby', 'Swift', 'Kotlin', 'Scala', 'R', 'MATLAB'
            ],
            'web_frameworks': [
                'React', 'Vue.js', 'Angular', 'Node.js', 'Express', 'Django', 'Flask',
                'Spring', 'Laravel', 'Symfony', 'ASP.NET', 'Next.js', 'Nuxt.js'
            ],
            'databases': [
                'PostgreSQL', 'MySQL', 'MongoDB', 'Redis', 'Oracle', 'SQL Server',
                'Cassandra', 'Elasticsearch', 'DynamoDB'
            ],
            'cloud': [
                'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Terraform',
                'CI/CD', 'Jenkins', 'GitLab CI', 'GitHub Actions'
            ],
            'tools': [
                'Git', 'Jira', 'Confluence', 'Slack', 'Agile', 'Scrum', 'DevOps'
            ],
        }

